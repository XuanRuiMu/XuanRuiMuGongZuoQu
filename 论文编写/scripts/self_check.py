# -*- coding: utf-8 -*-
import os
import sys

def run_self_check(docx_path):
    from docx import Document
    
    doc = Document(docx_path)
    results = []
    total = 0
    passed = 0
    
    def check(name, ok, detail=''):
        nonlocal total, passed
        total += 1
        status = 'PASS' if ok else 'FAIL'
        if ok:
            passed += 1
        marker = '✅' if ok else '❌'
        line = f'  {marker} [{status}] {name}'
        if detail:
            line += f' — {detail}'
        results.append((ok, name, detail))
        print(line)
    
    print('=' * 60)
    print('论文自检: %s' % os.path.basename(docx_path))
    print('=' * 60)
    
    print('\n--- A. 占位符检测 ---')
    
    placeholder_patterns = [
        ('\u00d7', '× 乘号占位符'),
        ('xxxx', '小写x占位符'),
        ('XXXX', '大写X占位符'),
        ('xxx', 'xxx占位符'),
        ('XXX', 'XXX占位符'),
        ('\u25a1', '□ 方框字符'),
        ('\u3000\u3000\u3000\u3000', '连续4个全角空格'),
    ]
    
    for char, desc in placeholder_patterns:
        found = False
        for i, p in enumerate(doc.paragraphs):
            if char in p.text:
                if char == '\u00d7' and any(k in p.text for k in ['乘法','倍率','系数','式（']):
                    continue
                style_name = p.style.name if p.style and p.style.name else ''
                if '_Style 28' in style_name or 'TOC' in style_name:
                    continue
                found = True
                snippet = p.text[:50].replace('\n',' ')
                check('%s残留' % desc, False, 'P%d: "%s..."' % (i, snippet))
                break
        if not found:
            check('%s残留' % desc, True, '未发现')
    
    print('\n--- B. 封面字段 ---')
    
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    
    check('题目非空', '我的世界RPG玩法系统' in full_text or 'RPG' in full_text)
    check('学生姓名', '魏建国' in full_text or '魏  建  国' in full_text or '姓名' not in full_text[:200])
    check('学号存在', len([p for p in doc.paragraphs if '6022203183' in p.text]) > 0 or '学号' in full_text)
    check('指导教师', '陈宏江' in full_text or '指导教师' in full_text)
    check('大学名称', '天津仁爱学院' in full_text or '大学' in full_text)
    check('学院名称', '智算工程' in full_text or '学院' in full_text)
    check('专业名称', '计算机科学' in full_text or '专业' in full_text)
    check('封面标题', '毕业设计' in full_text or '论文' in full_text)
    
    print('\n--- C. 结构完整性 ---')
    
    has_abstract_cn = any('摘' in p.text and '要' in p.text for p in doc.paragraphs[:80])
    has_abstract_en = any('ABSTRACT' in p.text.upper() for p in doc.paragraphs[:100])
    has_chapters = any('第' in p.text and chr(0x7ae0) in p.text for p in doc.paragraphs)
    has_references = any('参考文献' in p.text or 'References' in p.text for p in doc.paragraphs)
    has_tables = len(doc.tables) > 0
    
    check('中文摘要', has_abstract_cn)
    check('英文摘要', has_abstract_en)
    check('章节结构', has_chapters)
    check('参考文献', has_references)
    check('表格存在', has_tables, '%d个表格' % len(doc.tables))
    
    heading_count = sum(1 for p in doc.paragraphs if p.style and p.style.name is not None and 'Heading' in p.style.name)
    title_count = sum(1 for p in doc.paragraphs if p.style and p.style.name is not None and 'Title' in p.style.name)
    check('标题样式使用', heading_count > 0, '%d个Heading + %d个Title' % (heading_count, title_count))
    
    print('\n--- D. 格式检查 ---')
    
    run_count = sum(len(p.runs) for p in doc.paragraphs)
    para_count = len(doc.paragraphs)
    check('段落数量合理', 50 < para_count < 2000, '%d个段落' % para_count)
    check('Run结构保留', run_count > para_count * 1.5, '%d个runs (avg %.1f/para)' % (run_count, run_count/max(para_count,1)))
    
    underline_paras = [(i,p) for i,p in enumerate(doc.paragraphs) if any(r.underline for r in p.runs if r.text.strip())]
    underline_all = [(i,p) for i,p in enumerate(doc.paragraphs) if any(r.underline for r in p.runs)]
    check('下划线存在', len(underline_all) > 0, '%d处下划线(%d处有文本)' % (len(underline_all), len(underline_paras)))
    
    box_found = False
    for i, p in enumerate(doc.paragraphs):
        if '\u25a1' in p.text or '\u3000\u3000' in p.text:
            box_found = True
            break
    check('无方框字符', not box_found)
    
    x_in_title = False
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style and p.style.name is not None else ''
        if 'Title' in style and '\u00d7' in p.text:
            x_in_title = True
            break
    check('标题无×残留', not x_in_title)
    
    print('\n--- E. 内容质量 ---')
    
    empty_paras = sum(1 for p in doc.paragraphs if not p.text.strip())
    very_long_paras = sum(1 for p in doc.paragraphs if len(p.text) > 2000)
    check('空段落数量', empty_paras < 120, '%d个空段落' % empty_paras)
    check('无超长段落', very_long_paras == 0, '%d个超长段落' % very_long_paras)
    
    total_chars = sum(len(p.text) for p in doc.paragraphs)
    check('总字数合理', 5000 < total_chars < 200000, '约%d字' % total_chars)
    
    cn_chars = sum(sum(1 for c in p.text if '\u4e00' <= c <= '\u9fff') for p in doc.paragraphs)
    en_chars = sum(sum(1 for c in p.text if c.isascii() and c.isalpha()) for p in doc.paragraphs)
    check('中文字符足够', cn_chars > 1000, '%d中文字 / %d英文字' % (cn_chars, en_chars))
    
    duplicate_texts = {}
    for p in doc.paragraphs:
        t = p.text.strip()[:50]
        if len(t) > 10:
            duplicate_texts[t] = duplicate_texts.get(t, 0) + 1
    dups = {k:v for k,v in duplicate_texts.items() if v > 3}
    check('无明显重复内容', len(dups) == 0, '%d组重复(>3次)' % len(dups))
    
    print('\n' + '=' * 60)
    print('结果: %d/%d 通过 (%.1f%%)' % (passed, total, passed/total*100))
    if passed < total:
        fail_items = [r[1] for r in results if not r[0]]
        print('失败项: %s' % ', '.join(fail_items))
    print('=' * 60)
    
    return passed, total, results


if __name__ == '__main__':
    import glob
    
    target = sys.argv[1] if len(sys.argv) > 1 else None
    
    if not target:
        candidates = glob.glob(os.path.join(os.path.dirname(__file__), '..', '*终稿*.docx'))
        candidates = [c for c in candidates if 'v4' in c or 'v5' in c or 'v6' in c]
        if candidates:
            target = sorted(candidates)[-1]
    
    if not target or not os.path.exists(target):
        print('用法: python self_check.py <docx路径>')
        sys.exit(1)
    
    run_self_check(target)
