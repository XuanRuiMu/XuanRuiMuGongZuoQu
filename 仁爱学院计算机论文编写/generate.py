import os
import sys
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(SCRIPT_DIR, "thesis_template.docx")
DEFAULT_OUTPUT = os.path.join(SCRIPT_DIR, "..", "..", "..", "output.docx")
CONTENT_MD_PATH = os.path.join(SCRIPT_DIR, "..", "..", "..", "论文内容", "thesis_content.md")

_CODE_STYLES = {"代码——玄锐暮", "程序清单-附录——玄锐暮"}

_bookmark_id_counter = [0]


def _next_bookmark_id():
    _bookmark_id_counter[0] += 1
    return _bookmark_id_counter[0]


def _add_bookmark_to_para(para_elem, bookmark_name):
    bm_id = _next_bookmark_id()
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), str(bm_id))
    bm_start.set(qn('w:name'), bookmark_name)
    para_elem.insert(0, bm_start)
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(bm_id))
    para_elem.append(bm_end)


def _create_ref_field_runs(sdef, ref_num, display_text):
    runs = []
    r_begin = OxmlElement('w:r')
    rPr_begin = OxmlElement('w:rPr')
    if sdef.get('eastAsia_font') or sdef.get('ascii_font') or sdef.get('hAnsi_font'):
        rFonts = OxmlElement('w:rFonts')
        if sdef.get('eastAsia_font'):
            rFonts.set(qn('w:eastAsia'), sdef['eastAsia_font'])
        if sdef.get('ascii_font'):
            rFonts.set(qn('w:ascii'), sdef['ascii_font'])
        if sdef.get('hAnsi_font'):
            rFonts.set(qn('w:hAnsi'), sdef['hAnsi_font'])
        rPr_begin.append(rFonts)
    if sdef.get('font_size') is not None:
        half_pt = _emu_to_half_pt(sdef['font_size'])
        if half_pt is not None:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(half_pt))
            rPr_begin.append(sz)
    if rPr_begin.findall('*'):
        r_begin.append(rPr_begin)
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    r_begin.append(fldChar_begin)
    runs.append(r_begin)

    r_instr = OxmlElement('w:r')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' REF ref_{ref_num} \\h '
    r_instr.append(instrText)
    runs.append(r_instr)

    r_sep = OxmlElement('w:r')
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    r_sep.append(fldChar_sep)
    runs.append(r_sep)

    r_display = _create_run_element(sdef, display_text, superscript=True, convert_quotes=False)
    runs.append(r_display)

    r_end = OxmlElement('w:r')
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    r_end.append(fldChar_end)
    runs.append(r_end)

    return runs


def _convert_quotes_to_chinese(text):
    """Convert ASCII double quotes to Chinese double quotes.
    Uses a smarter approach: reads quote pairs from left to right.
    If unmatched quotes remain (odd count), leaves the last one as-is
    to avoid corrupting the text."""
    result = []
    in_quote = False
    quote_positions = [i for i, ch in enumerate(text) if ch == '"']
    # If odd number, treat the last one as unmatched (likely part of code or malformed)
    valid_pairs = len(quote_positions) // 2 * 2
    pair_set = set(quote_positions[:valid_pairs])
    pair_idx = 0

    for i, ch in enumerate(text):
        if ch == '"' and i in pair_set:
            if not in_quote:
                result.append('\u201c')
                in_quote = True
            else:
                result.append('\u201d')
                in_quote = False
        else:
            result.append(ch)
    return ''.join(result)


def _convert_en_abstract_punctuation(text):
    parts = re.split(r'(https?://\S+)', text)
    result = []
    for i, part in enumerate(parts):
        if i % 2 == 1:
            result.append(part)
        else:
            part = part.replace(':', '：')
            part = part.replace(';', '；')
            result.append(part)
    return ''.join(result)


def _make_run(text, eastAsia=None, ascii_f=None, hAnsi_f=None):
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    has = False
    if eastAsia or ascii_f or hAnsi_f:
        rFonts = OxmlElement('w:rFonts')
        if eastAsia:
            rFonts.set(qn('w:eastAsia'), eastAsia)
        if ascii_f:
            rFonts.set(qn('w:ascii'), ascii_f)
            if not hAnsi_f:
                rFonts.set(qn('w:hAnsi'), ascii_f)
        if hAnsi_f:
            rFonts.set(qn('w:hAnsi'), hAnsi_f)
        rPr.append(rFonts)
        has = True
    if has:
        r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    return r


ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: 'left',
    WD_ALIGN_PARAGRAPH.CENTER: 'center',
    WD_ALIGN_PARAGRAPH.RIGHT: 'right',
    WD_ALIGN_PARAGRAPH.JUSTIFY: 'both',
    None: None,
}


def read_style_defs(doc):
    defs = {}
    for style in doc.styles:
        if style.name and '玄锐暮' in style.name:
            sdef = {'style_id': style.style_id, 'name': style.name}
            pf = style.paragraph_format
            if pf:
                sdef['alignment'] = ALIGN_MAP.get(pf.alignment, None)
                sdef['first_line_indent'] = pf.first_line_indent
                sdef['left_indent'] = pf.left_indent
                sdef['line_spacing'] = pf.line_spacing
                sdef['line_spacing_rule'] = pf.line_spacing_rule
                sdef['space_before'] = pf.space_before
                sdef['space_after'] = pf.space_after
            font = style.font
            if font:
                sdef['font_name'] = font.name
                sdef['font_size'] = font.size
                sdef['font_bold'] = font.bold
            rPr = style.element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
            if rPr is not None:
                rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
                if rFonts is not None:
                    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                    sdef['eastAsia_font'] = rFonts.get(ns + 'eastAsia')
                    sdef['ascii_font'] = rFonts.get(ns + 'ascii')
                    sdef['hAnsi_font'] = rFonts.get(ns + 'hAnsi')
            defs[style.name] = sdef
    return defs


def _emu_to_twip(emu):
    if emu is None:
        return None
    return int(emu / 12700 * 20)


def _emu_to_half_pt(emu):
    if emu is None:
        return None
    return int(emu / 12700 * 2)


def _create_run_element(sdef, text, superscript=False, convert_quotes=True):
    r = OxmlElement('w:r')
    rPr_elem = OxmlElement('w:rPr')
    has_rPr = False
    if sdef.get('eastAsia_font') or sdef.get('ascii_font') or sdef.get('hAnsi_font'):
        rFonts = OxmlElement('w:rFonts')
        if sdef.get('eastAsia_font'):
            rFonts.set(qn('w:eastAsia'), sdef['eastAsia_font'])
        if sdef.get('ascii_font'):
            rFonts.set(qn('w:ascii'), sdef['ascii_font'])
        if sdef.get('hAnsi_font'):
            rFonts.set(qn('w:hAnsi'), sdef['hAnsi_font'])
        rPr_elem.append(rFonts)
        has_rPr = True
    if sdef.get('font_size') is not None:
        half_pt = _emu_to_half_pt(sdef['font_size'])
        if half_pt is not None:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(half_pt))
            rPr_elem.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(half_pt))
            rPr_elem.append(szCs)
            has_rPr = True
    if sdef.get('font_bold') is True:
        b = OxmlElement('w:b')
        rPr_elem.append(b)
        bCs = OxmlElement('w:bCs')
        rPr_elem.append(bCs)
        has_rPr = True
    elif sdef.get('font_bold') is False:
        b = OxmlElement('w:b')
        b.set(qn('w:val'), 'false')
        rPr_elem.append(b)
        bCs = OxmlElement('w:bCs')
        bCs.set(qn('w:val'), 'false')
        rPr_elem.append(bCs)
        has_rPr = True
    if superscript:
        vertAlign = OxmlElement('w:vertAlign')
        vertAlign.set(qn('w:val'), 'superscript')
        rPr_elem.append(vertAlign)
        has_rPr = True
    if has_rPr:
        r.append(rPr_elem)
    t = OxmlElement('w:t')
    if convert_quotes:
        text = _convert_quotes_to_chinese(text)
    text = text.replace("`", "")  # 移除markdown反引号
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    return r


def create_para_element(style_name, text, style_defs, page_break=False):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    sdef = style_defs.get(style_name, {})
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), sdef.get('style_id', style_name))
    pPr.append(pStyle)

    if sdef.get('alignment'):
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), sdef['alignment'])
        pPr.append(jc)

    if sdef.get('first_line_indent') is not None:
        ind = OxmlElement('w:ind')
        twip = _emu_to_twip(sdef['first_line_indent'])
        if twip is not None:
            ind.set(qn('w:firstLine'), str(twip))
        if sdef.get('left_indent') is not None:
            left_twip = _emu_to_twip(sdef['left_indent'])
            if left_twip is not None:
                ind.set(qn('w:left'), str(left_twip))
        pPr.append(ind)
    elif sdef.get('left_indent') is not None:
        ind = OxmlElement('w:ind')
        left_twip = _emu_to_twip(sdef['left_indent'])
        if left_twip is not None:
            ind.set(qn('w:left'), str(left_twip))
        pPr.append(ind)

    if sdef.get('line_spacing') is not None or sdef.get('space_before') is not None or sdef.get('space_after') is not None:
        spacing = OxmlElement('w:spacing')
        if sdef.get('line_spacing') is not None:
            if sdef.get('line_spacing_rule') == WD_LINE_SPACING.MULTIPLE:
                spacing.set(qn('w:line'), str(int(sdef['line_spacing'] * 240)))
                spacing.set(qn('w:lineRule'), 'auto')
            else:
                twip = _emu_to_twip(sdef['line_spacing'])
                if twip is not None:
                    spacing.set(qn('w:line'), str(twip))
                    spacing.set(qn('w:lineRule'), 'exact')
        if sdef.get('space_before') is not None:
            twip = _emu_to_twip(sdef['space_before'])
            if twip is not None:
                spacing.set(qn('w:before'), str(twip))
        if sdef.get('space_after') is not None:
            twip = _emu_to_twip(sdef['space_after'])
            if twip is not None:
                spacing.set(qn('w:after'), str(twip))
        pPr.append(spacing)

    if page_break:
        pb = OxmlElement('w:pageBreakBefore')
        pPr.append(pb)

    p.append(pPr)

    is_code = style_name in _CODE_STYLES
    if is_code:
        r = _create_run_element(sdef, text, superscript=False, convert_quotes=False)
        p.append(r)
    else:
        text = _convert_quotes_to_chinese(text)
        text = text.replace("`", "")  # 移除markdown反引号
        parts = re.split(r'(\[\d+\])', text)
        for part in parts:
            if not part:
                continue
            if re.match(r'^\[\d+\]$', part):
                ref_num = part[1:-1]
                ref_runs = _create_ref_field_runs(sdef, ref_num, part)
                for rr in ref_runs:
                    p.append(rr)
            else:
                r = _create_run_element(sdef, part, superscript=False, convert_quotes=False)
                p.append(r)

    return p


def create_page_break_element():
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    return p


def merge_runs_to_first(paragraph):
    if not paragraph.runs:
        return ""
    full = "".join(r.text for r in paragraph.runs)
    paragraph.runs[0].text = full
    for r in paragraph.runs[1:]:
        r.text = ""
    return full


def set_cell_text(cell, text, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.style = cell.part.document.styles['表格内容——玄锐暮']
    p.alignment = align
    run = p.add_run(text)


def _set_cell_border(cell, top=None, bottom=None):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    old = tcPr.find(qn('w:tcBorders'))
    if old is not None:
        tcPr.remove(old)
    borders = OxmlElement('w:tcBorders')
    for name, val in [('top', top), ('bottom', bottom)]:
        if val is None:
            continue
        b = OxmlElement(f'w:{name}')
        if val == 'nil':
            b.set(qn('w:val'), 'nil')
        else:
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), str(val))
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), '000000')
        borders.append(b)
    tcPr.append(borders)


def _set_row_pr_ex(row):
    tr = row._tr
    tblPrEx = OxmlElement('w:tblPrEx')
    tblBorders = OxmlElement('w:tblBorders')
    for name, val, sz in [
        ('top', 'single', '12'), ('left', 'none', '0'),
        ('bottom', 'single', '12'), ('right', 'none', '0'),
        ('insideH', 'single', '4'), ('insideV', 'none', '0'),
    ]:
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), val)
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        tblBorders.append(b)
    tblPrEx.append(tblBorders)
    tr.insert(0, tblPrEx)
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(1, trPr)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    trPr.append(jc)


def make_three_line_table(doc, headers, rows):
    num_cols = len(headers)
    total_rows = 1 + len(rows)
    tbl = doc.add_table(rows=total_rows, cols=num_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(tbl.rows[0].cells[i], h)
    for ri, row_data in enumerate(rows):
        for ci in range(num_cols):
            val = row_data[ci] if ci < len(row_data) else ""
            set_cell_text(tbl.rows[ri + 1].cells[ci], str(val))
    tbl_element = tbl._tbl
    tblPr = tbl_element.tblPr if tbl_element.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    top_border = OxmlElement('w:top')
    top_border.set(qn('w:val'), 'single')
    top_border.set(qn('w:sz'), '12')
    top_border.set(qn('w:space'), '0')
    top_border.set(qn('w:color'), '000000')
    borders.append(top_border)
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '12')
    bottom_border.set(qn('w:space'), '0')
    bottom_border.set(qn('w:color'), '000000')
    borders.append(bottom_border)
    for border_name in ['left', 'right', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'single')
    insideH.set(qn('w:sz'), '4')
    insideH.set(qn('w:space'), '0')
    insideH.set(qn('w:color'), '000000')
    borders.append(insideH)
    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    tblPr.append(borders)
    for ci in range(num_cols):
        _set_cell_border(tbl.rows[0].cells[ci], bottom=4)
    for ri in range(1, total_rows):
        _set_row_pr_ex(tbl.rows[ri])
        for ci in range(num_cols):
            if ri == 1:
                _set_cell_border(tbl.rows[ri].cells[ci], top=4, bottom='nil')
            elif ri < total_rows - 1:
                _set_cell_border(tbl.rows[ri].cells[ci], top='nil', bottom='nil')
            else:
                _set_cell_border(tbl.rows[ri].cells[ci], top='nil')
    return tbl


def process_block(style_defs, block, elements):
    btype = block.get("type", "para")
    if btype == "para":
        elem = create_para_element("正文——玄锐暮", block["text"], style_defs)
        elements.append(elem)
    elif btype == "code":
        elements.append(create_para_element("代码——玄锐暮", "", style_defs))
        for line in block["lines"]:
            elem = create_para_element("代码——玄锐暮", line, style_defs)
            elements.append(elem)
        elements.append(create_para_element("代码——玄锐暮", "", style_defs))
    elif btype == "figure":
        fig_caption = block.get("caption", "图X-X  图片")
        elem = create_para_element("表格、图片标题——玄锐暮", fig_caption, style_defs)
        elements.append(elem)
    elif btype == "table":
        elements.append(create_para_element("表格、图片标题——玄锐暮", "", style_defs))
        cap_elem = create_para_element("表格、图片标题——玄锐暮", block.get("caption", "表X-X  表格"), style_defs)
        elements.append(cap_elem)
        return ("table", block, elements)
    return None


def build_chapter_elements(style_defs, ch_data, page_break=True):
    elements = []
    ch_title = ch_data["title"]
    m = re.match(r'^(第.+章)\s*(.*)', ch_title)
    if m and m.group(2):
        ch_title = m.group(1) + '  ' + m.group(2)
    title_elem = create_para_element("一级标题——玄锐暮", ch_title, style_defs, page_break=page_break)
    elements.append(title_elem)

    for sec in ch_data.get("sections", []):
        if sec["heading2"] != ch_data["title"]:
            h2_elem = create_para_element("二级标题及以下——玄锐暮", sec["heading2"], style_defs)
            elements.append(h2_elem)

        for h3_data in sec.get("heading3_list", []):
            h3_elem = create_para_element("二级标题及以下——玄锐暮", h3_data["heading3"], style_defs)
            elements.append(h3_elem)

            for block in h3_data.get("blocks", []):
                result = process_block(style_defs, block, elements)
                if result and result[0] == "table":
                    elements.append(("__table__", result[1]))
                    elements.append(("__table_end__",))

        for block in sec.get("blocks", []):
            result = process_block(style_defs, block, elements)
            if result and result[0] == "table":
                elements.append(("__table__", result[1]))
                elements.append(("__table_end__",))

    return elements


def parse_markdown(md_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        raw = f.read()

    lines = raw.split('\n')
    body_start = 0

    if lines and lines[0].strip() == '---':
        for i in range(1, len(lines)):
            if lines[i].strip() == '---':
                body_start = i + 1
                break

    body_lines = lines[body_start:]

    result = {
        "abstract_cn": "",
        "keywords_cn": "",
        "abstract_en": "",
        "keywords_en": "",
        "chapters": [],
        "references": [],
        "acknowledgment": [],
        "appendix": [],
        "foreign_original": {
            "title": "",
            "author_affiliation": "",
            "abstract": "",
            "keywords": "",
            "body": [],
        },
        "foreign_translation": {
            "title": "",
            "author_affiliation": "",
            "abstract": "",
            "keywords": "",
            "body": [],
        },
    }

    section_markers = []
    for i, line in enumerate(body_lines):
        stripped = line.strip()
        if re.match(r'^# ', stripped):
            section_markers.append((i, stripped))

    abstract_idx = None
    abstract_en_idx = None
    chapter1_idx = None
    ref_idx = None
    ack_idx = None
    foreign_idx = None
    appendix_idx = None

    for idx, text in section_markers:
        if text == '# 摘要':
            abstract_idx = idx
        elif text == '# ABSTRACT':
            abstract_en_idx = idx
        elif re.match(r'^# 第.+章', text) or re.match(r'^# 第.+章\s', text):
            if chapter1_idx is None:
                chapter1_idx = idx
        elif text == '# 参考文献':
            ref_idx = idx
        elif text == '# 致谢':
            ack_idx = idx
        elif text == '# 外文原文及译文':
            foreign_idx = idx
        elif text == '# 附录':
            appendix_idx = idx

    if abstract_idx is not None:
        abs_end = abstract_en_idx if abstract_en_idx is not None else (chapter1_idx or len(body_lines))
        abs_lines = [body_lines[i].rstrip() for i in range(abstract_idx + 1, abs_end)
                     if body_lines[i].strip()]
        if abs_lines:
            kw_line = None
            for li in reversed(range(len(abs_lines))):
                if abs_lines[li].startswith('关键词') or abs_lines[li].startswith('关键词：') or abs_lines[li].startswith('关键词:'):
                    kw_line = abs_lines[li]
                    abs_lines = abs_lines[:li]
                    break
            result["abstract_cn"] = '\n'.join(abs_lines)
            if kw_line:
                kw_match = re.search(r'关键词[：:]\s*(.+)', kw_line)
                if kw_match:
                    result["keywords_cn"] = kw_match.group(1).strip()

    if abstract_en_idx is not None:
        aben_end = chapter1_idx if chapter1_idx is not None else (ref_idx or len(body_lines))
        aben_lines = [body_lines[i].rstrip() for i in range(abstract_en_idx + 1, aben_end)
                      if body_lines[i].strip()]
        if aben_lines:
            kw_line = None
            for li in reversed(range(len(aben_lines))):
                if aben_lines[li].startswith('Key words') or aben_lines[li].startswith('Key words:') or aben_lines[li].startswith('Key words：'):
                    kw_line = aben_lines[li]
                    aben_lines = aben_lines[:li]
                    break
            result["abstract_en"] = '\n'.join(aben_lines)
            if kw_line:
                kw_match = re.search(r'Key words[：:]\s*(.+)', kw_line)
                if kw_match:
                    result["keywords_en"] = kw_match.group(1).strip()

    chapter_indices = []
    for idx, text in section_markers:
        if re.match(r'^# 第.+章', text) or re.match(r'^# 第.+章\s', text):
            chapter_indices.append((idx, text))

    for ci, (ch_start, ch_title_text) in enumerate(chapter_indices):
        ch_end_idx = chapter_indices[ci + 1][0] if ci + 1 < len(chapter_indices) else (ref_idx or len(body_lines))
        ch_raw_lines = body_lines[ch_start + 1:ch_end_idx]

        ch_title = re.sub(r'^#\s+', '', ch_title_text).strip()
        is_first = (ci == 0)

        sections = []
        current_h2 = None
        current_h3 = None

        i = 0
        while i < len(ch_raw_lines):
            line = ch_raw_lines[i].rstrip()
            stripped = line.strip()

            h3_match = re.match(r'^###\s+(.+)$', stripped)
            if h3_match:
                current_h3 = {"heading3": h3_match.group(1).strip(), "blocks": []}
                if current_h2 is not None:
                    if "heading3_list" not in current_h2:
                        current_h2["heading3_list"] = []
                    current_h2["heading3_list"].append(current_h3)
                i += 1
                continue

            h2_match = re.match(r'^##\s+(.+)$', stripped)
            if h2_match:
                current_h2 = {"heading2": h2_match.group(1).strip(), "blocks": [], "heading3_list": []}
                sections.append(current_h2)
                current_h3 = None
                i += 1
                continue

            fig_match = re.match(r'^!\[(.+?)\]\((.+?)\)$', stripped)
            if fig_match:
                block = {"type": "figure", "caption": fig_match.group(1).strip()}
                target = current_h3 if current_h3 is not None else current_h2
                if target is None:
                    if not sections:
                        sections.append({"heading2": ch_title, "blocks": [], "heading3_list": []})
                    target = sections[-1]
                target["blocks"].append(block)
                i += 1
                continue

            code_match = re.match(r'^```\w*$', stripped)
            if code_match:
                code_lines = []
                i += 1
                while i < len(ch_raw_lines) and not re.match(r'^```$', ch_raw_lines[i].strip()):
                    code_line = re.sub(r'^\s*\d+\s{2,}', '', ch_raw_lines[i].rstrip())
                    code_lines.append(code_line)
                    i += 1
                i += 1
                block = {"type": "code", "lines": code_lines}
                target = current_h3 if current_h3 is not None else current_h2
                if target is None:
                    if not sections:
                        sections.append({"heading2": ch_title, "blocks": [], "heading3_list": []})
                    target = sections[-1]
                target["blocks"].append(block)
                continue

            table_match = re.match(r'^\|(.+)\|$', stripped)
            if table_match:
                table_rows = []
                while i < len(ch_raw_lines) and re.match(r'^\|', ch_raw_lines[i].strip()):
                    trow = ch_raw_lines[i].strip()
                    if re.match(r'^\|[\s\-:|]+\|$', trow):
                        i += 1
                        continue
                    cells = [c.strip() for c in trow.strip('|').split('|')]
                    table_rows.append(cells)
                    i += 1
                # 过滤全空行（所有单元格均为空的行）
                table_rows = [row for row in table_rows if any(c.strip() for c in row)]
                if len(table_rows) >= 2:
                    headers = table_rows[0]
                    data_rows = table_rows[1:]
                    caption_text = None
                    target_obj = current_h3 if current_h3 is not None else current_h2
                    if target_obj is None:
                        if not sections:
                            sections.append({"heading2": ch_title, "blocks": [], "heading3_list": []})
                        target_obj = sections[-1]
                    tbl_caption_pat = re.compile(r'^表\d+-\d+\s+.+$')
                    if target_obj and "blocks" in target_obj and len(target_obj["blocks"]) > 0:
                        scan_idx = len(target_obj["blocks"]) - 1
                        while scan_idx >= 0:
                            blk = target_obj["blocks"][scan_idx]
                            if blk.get("type") == "para" and blk.get("text", "").strip() == "":
                                scan_idx -= 1
                                continue
                            if blk.get("type") == "para" and tbl_caption_pat.match(blk.get("text", "")):
                                caption_text = blk["text"]
                                del target_obj["blocks"][scan_idx:]
                                break
                            break
                    if caption_text is None:
                        raw_heading = current_h2.get('heading2', '表格') if current_h2 else ch_title
                        clean_heading = re.sub(r'^\d+(\.\d+)*\s+', '', raw_heading)
                        caption_text = f"表X-X  {clean_heading}"
                    block = {"type": "table", "caption": caption_text, "headers": headers, "rows": data_rows}
                    target_obj["blocks"].append(block)
                continue

            if stripped and not stripped.startswith('#'):
                block = {"type": "para", "text": stripped}
                target = current_h3 if current_h3 is not None else current_h2
                if target is None:
                    if not sections:
                        sections.append({"heading2": ch_title, "blocks": [], "heading3_list": []})
                    target = sections[-1]
                target["blocks"].append(block)

            i += 1

        result["chapters"].append({
            "title": ch_title,
            "is_first_chapter": is_first,
            "sections": sections,
        })

    if ref_idx is not None:
        ref_end = len(body_lines)
        for end_candidate in [ack_idx, appendix_idx, foreign_idx]:
            if end_candidate is not None and end_candidate < ref_end:
                ref_end = end_candidate
        for i in range(ref_idx + 1, ref_end):
            stripped = body_lines[i].strip()
            if stripped and re.match(r'^\[[0-9]+\]', stripped):
                result["references"].append(stripped)

    if ack_idx is not None:
        ack_end = len(body_lines)
        for end_candidate in [appendix_idx, foreign_idx]:
            if end_candidate is not None and end_candidate > ack_idx and end_candidate < ack_end:
                ack_end = end_candidate
        for i in range(ack_idx + 1, ack_end):
            stripped = body_lines[i].strip()
            if stripped:
                result["acknowledgment"].append(stripped)

    if foreign_idx is not None:
        foreign_end = len(body_lines)
        for marker_idx, marker_text in section_markers:
            if marker_idx > foreign_idx:
                foreign_end = marker_idx
                break

        in_translation = False
        current_subsection = None
        in_body = False

        for i in range(foreign_idx + 1, foreign_end):
            stripped = body_lines[i].strip()
            if not stripped:
                continue

            if stripped == '## 中文译文':
                in_translation = True
                current_subsection = None
                in_body = False
                continue
            if stripped == '## 外文原文':
                continue

            sub_match = re.match(r'^###\s+\[(.+?)\]$', stripped)
            if sub_match:
                current_subsection = sub_match.group(1).strip()
                in_body = (current_subsection == '正文')
                continue

            h4_match = re.match(r'^#{4,5}\s+(.+)$', stripped)
            if h4_match and in_body:
                target = result["foreign_translation"] if in_translation else result["foreign_original"]
                target["body"].append({"type": "heading", "text": h4_match.group(1).strip()})
                continue

            if in_body:
                target = result["foreign_translation"] if in_translation else result["foreign_original"]
                target["body"].append({"type": "para", "text": stripped})
                continue

            target = result["foreign_translation"] if in_translation else result["foreign_original"]
            if current_subsection == '标题':
                target["title"] = stripped
            elif current_subsection in ('作者、单位', '作者,单位'):
                if target["author_affiliation"]:
                    target["author_affiliation"] += "\n" + stripped
                else:
                    target["author_affiliation"] = stripped
            elif current_subsection == '摘要':
                if target["abstract"]:
                    target["abstract"] += " " + stripped
                else:
                    target["abstract"] = stripped
            elif current_subsection == '关键词':
                target["keywords"] = stripped

    if appendix_idx is not None:
        appendix_end = len(body_lines)
        for marker_idx, marker_text in section_markers:
            if marker_idx > appendix_idx:
                appendix_end = marker_idx
                break
        in_code = False
        for i in range(appendix_idx + 1, appendix_end):
            stripped = body_lines[i].strip()
            if re.match(r'^```\w*$', stripped):
                in_code = not in_code
                continue
            if in_code:
                result["appendix"].append(body_lines[i].rstrip())

    return result


def _fix_reference_numbering(result):
    references = result.get("references", [])
    if not references:
        return result

    ref_nums = []
    for ref in references:
        m = re.match(r'^\[(\d+)\]', ref)
        if m:
            ref_nums.append(int(m.group(1)))

    if not ref_nums:
        return result

    expected = list(range(1, len(ref_nums) + 1))
    if ref_nums == expected:
        return result

    old_to_new = {}
    for new_num, old_num in enumerate(sorted(ref_nums), 1):
        old_to_new[old_num] = new_num

    new_references = []
    for ref in references:
        m = re.match(r'^\[(\d+)\]', ref)
        if m:
            old_num = int(m.group(1))
            new_num = old_to_new.get(old_num, old_num)
            new_ref = re.sub(r'^\[\d+\]', f'[{new_num}]', ref)
            new_references.append(new_ref)
        else:
            new_references.append(ref)
    result["references"] = new_references

    def _renumber_citations(text):
        def _replace_citation(m):
            old_num = int(m.group(1))
            new_num = old_to_new.get(old_num, old_num)
            return f'[{new_num}]'
        return re.sub(r'\[(\d+)\]', _replace_citation, text)

    def _fix_blocks(blocks):
        for block in blocks:
            if block.get("type") == "para" and block.get("text"):
                block["text"] = _renumber_citations(block["text"])
            elif block.get("type") == "table":
                if block.get("caption"):
                    block["caption"] = _renumber_citations(block["caption"])
                for row in block.get("rows", []):
                    for i, cell in enumerate(row):
                        row[i] = _renumber_citations(cell)

    for chapter in result.get("chapters", []):
        for section in chapter.get("sections", []):
            _fix_blocks(section.get("blocks", []))
            for h3 in section.get("heading3_list", []):
                _fix_blocks(h3.get("blocks", []))

    for key in ["abstract_cn", "abstract_en"]:
        if result.get(key):
            result[key] = _renumber_citations(result[key])

    for key in ["acknowledgment"]:
        new_ack = []
        for line in result.get(key, []):
            new_ack.append(_renumber_citations(line))
        result[key] = new_ack

    print(f"  ✓ 参考文献编号修复: {ref_nums} → {expected}")
    return result


def phase1_render_template(content):
    print("阶段1: 渲染模板 (摘要/关键词)...")
    doc = Document(TEMPLATE_PATH)

    abstract_cn = _convert_quotes_to_chinese(content.get("abstract_cn", ""))
    keywords_cn = content.get("keywords_cn", "")
    abstract_en = _convert_quotes_to_chinese(content.get("abstract_en", ""))
    keywords_en = content.get("keywords_en", "")

    abstract_cn_idx = None
    abstract_en_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "摘    要" and abstract_cn_idx is None:
            abstract_cn_idx = i
        elif p.text.strip() == "ABSTRACT" and abstract_en_idx is None:
            abstract_en_idx = i

    style_defs_local = read_style_defs(doc)

    cn_abs_para = None
    cn_kw_para = None
    en_abs_para = None
    en_kw_para = None
    for i, p in enumerate(doc.paragraphs):
        if abstract_cn_idx is not None and abstract_en_idx is not None:
            if abstract_cn_idx < i < abstract_en_idx:
                if len(p.text) > 100 and cn_abs_para is None:
                    cn_abs_para = p
                elif p.text.strip().startswith("关键词") and cn_kw_para is None:
                    cn_kw_para = p
            elif i > abstract_en_idx:
                if len(p.text) > 100 and en_abs_para is None and not p.text.strip().startswith("Key"):
                    if not any('\u4e00' <= c <= '\u9fff' for c in p.text[:20]):
                        en_abs_para = p
                elif p.text.strip().startswith("Key words") and en_kw_para is None:
                    en_kw_para = p

    if cn_abs_para is not None and abstract_cn:
        cn_abs_paragraphs = [line for line in abstract_cn.split('\n') if line.strip()]
        if cn_abs_paragraphs:
            merge_runs_to_first(cn_abs_para)
            if cn_abs_para.runs:
                cn_abs_para.runs[0].text = cn_abs_paragraphs[0]
            for r in cn_abs_para.runs[1:]:
                r.text = ""
            prev_elem = cn_abs_para._element
            abs_cn_style = "正文——玄锐暮"  # 摘要正文使用正文样式
            for para_text in cn_abs_paragraphs[1:]:
                elem = create_para_element(abs_cn_style, para_text, style_defs_local)
                prev_elem.addnext(elem)
                prev_elem = elem

    if cn_kw_para is not None and keywords_cn:
        p_elem = cn_kw_para._element
        for r_elem in list(p_elem.findall(qn('w:r'))):
            p_elem.remove(r_elem)
        p_elem.append(_make_run("关键词：", eastAsia="黑体", ascii_f="黑体"))
        p_elem.append(_make_run(keywords_cn, eastAsia="宋体", ascii_f="Times New Roman"))

    if en_abs_para is not None and abstract_en:
        abstract_en = _convert_en_abstract_punctuation(abstract_en)
        en_abs_paragraphs = [line for line in abstract_en.split('\n') if line.strip()]
        if en_abs_paragraphs:
            merge_runs_to_first(en_abs_para)
            if en_abs_para.runs:
                en_abs_para.runs[0].text = en_abs_paragraphs[0]
            for r in en_abs_para.runs[1:]:
                r.text = ""
            prev_elem = en_abs_para._element
            abs_en_style = "正文——玄锐暮"  # 英文摘要正文使用正文样式
            for para_text in en_abs_paragraphs[1:]:
                para_text = _convert_en_abstract_punctuation(para_text)
                elem = create_para_element(abs_en_style, para_text, style_defs_local)
                prev_elem.addnext(elem)
                prev_elem = elem

    if en_kw_para is not None and keywords_en:
        keywords_en = _convert_en_abstract_punctuation(keywords_en)
        p_elem = en_kw_para._element
        for r_elem in list(p_elem.findall(qn('w:r'))):
            p_elem.remove(r_elem)
        p_elem.append(_make_run("Key words：", eastAsia="黑体", ascii_f="黑体"))
        p_elem.append(_make_run(keywords_en, eastAsia="宋体", ascii_f="Times New Roman"))

    temp_path = os.path.join(SCRIPT_DIR, "_temp_rendered.docx")
    doc.save(temp_path)
    print("  ✓ 模板渲染完成")
    return temp_path


def phase2_add_chapters(temp_path, content, style_defs):
    print("阶段2: 添加正文内容...")
    doc = Document(temp_path)

    ch1_title_elem = None
    ref_title_elem = None

    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        if style_name == "一级标题——玄锐暮" and "第一章" in p.text:
            ch1_title_elem = p._element
        elif style_name == "Title" and "参考文献" in p.text:
            ref_title_elem = p._element

    if ch1_title_elem is None:
        for p in doc.paragraphs:
            style_name = p.style.name if p.style else ""
            if style_name != "目录——玄锐暮" and "第一章" in p.text:
                ch1_title_elem = p._element
                break

    if ref_title_elem is None:
        for p in doc.paragraphs:
            style_name = p.style.name if p.style else ""
            if style_name != "目录——玄锐暮" and "参考文献" in p.text:
                ref_title_elem = p._element
                break

    ch1_prev_elem = None
    prev = None
    for child in list(doc.element.body):
        if child is ch1_title_elem:
            ch1_prev_elem = prev
        prev = child

    if ch1_title_elem is not None and ref_title_elem is not None:
        body = doc.element.body
        removing = False
        to_remove = []
        for child in list(body):
            if child is ch1_title_elem:
                to_remove.append(child)
                removing = True
                continue
            if child is ref_title_elem:
                removing = False
                break
            if removing:
                has_sectPr = child.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
                if has_sectPr is not None:
                    print(f"  ⚠ 跳过含sectPr的元素，避免破坏分节符")
                    continue
                to_remove.append(child)
        for elem in to_remove:
            body.remove(elem)
        print(f"  ✓ 清除旧正文内容: {len(to_remove)}个元素")
    else:
        print("  ⚠ 无法定位正文区域，跳过清除")

    if ch1_prev_elem is not None:
        total_elements = 0
        chapters = content.get("chapters", [])
        for ch in reversed(chapters):
            is_first = ch.get("is_first_chapter", False)
            ch_elements = build_chapter_elements(style_defs, ch, page_break=not is_first)
            for elem in reversed(ch_elements):
                if isinstance(elem, tuple) and elem[0] == "__table__":
                    tbl = make_three_line_table(doc, elem[1]["headers"], elem[1]["rows"])
                    ch1_prev_elem.addnext(tbl._element)
                elif isinstance(elem, tuple) and elem[0] == "__table_end__":
                    empty_elem = create_para_element("表格、图片标题——玄锐暮", "", style_defs)
                    ch1_prev_elem.addnext(empty_elem)
                else:
                    ch1_prev_elem.addnext(elem)
            total_elements += len(ch_elements)
        print(f"  ✓ 插入正文内容: {total_elements}个元素")
    else:
        print("  ⚠ 无法找到插入位置")

    return doc


def phase3_handle_references(doc, content, style_defs):
    print("阶段3: 处理参考文献...")
    references = content.get("references", [])
    if not references:
        print("  ⚠ 无参考文献数据")
        return doc

    ref_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.style and p.style.name == "Title" and "参考文献" in p.text:
            ref_idx = i
            break
    if ref_idx is None:
        for i, p in enumerate(doc.paragraphs):
            if "参考文献" in p.text and i > 50:
                ref_idx = i
                break

    if ref_idx is None:
        print("  ⚠ 未找到参考文献标题")
        return doc

    ref_para_indices = []
    for i in range(ref_idx + 1, len(doc.paragraphs)):
        pt = doc.paragraphs[i].text.strip()
        if pt.startswith("["):
            ref_para_indices.append(i)
        elif pt.startswith("【") or (pt.startswith("（") and "字体" in pt):
            break

    # 参考文献前插入分页符
    ref_heading = doc.paragraphs[ref_idx]
    pb = create_page_break_element()
    ref_heading._element.addprevious(pb)

    ref_style_name = "参考文献1~9——玄锐暮"
    for j, ri in enumerate(ref_para_indices):
        if j < len(references):
            p = doc.paragraphs[ri]
            merge_runs_to_first(p)
            if p.runs:
                p.runs[0].text = references[j]
                if j >= 9:
                    ref_style_name = "参考文献10~之后——玄锐暮"
                p.style = doc.styles[ref_style_name]
            for r in p.runs[1:]:
                r.text = ""
            _add_bookmark_to_para(p._element, f'ref_{j + 1}')

    body = doc.element.body
    if len(references) > len(ref_para_indices) and len(ref_para_indices) > 0:
        last_ref_elem = doc.paragraphs[ref_para_indices[-1]]._element
        for j in range(len(ref_para_indices), len(references)):
            ref_sname = "参考文献10~之后——玄锐暮" if j >= 9 else "参考文献1~9——玄锐暮"
            elem = create_para_element(ref_sname, references[j], style_defs)
            _add_bookmark_to_para(elem, f'ref_{j + 1}')
            last_ref_elem.addnext(elem)
            last_ref_elem = elem

    to_clean = []
    last_ref_idx = ref_para_indices[-1] if ref_para_indices else ref_idx
    for ci in range(last_ref_idx + 1, len(doc.paragraphs)):
        pp = doc.paragraphs[ci]
        t = pp.text.strip()
        if '附    录' in t or ('附' in t and '录' in t and pp.style.name == 'Title'):
            break
        if '致    谢' in t:
            break
        has_sectPr = pp._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
        if has_sectPr is not None:
            continue
        if t.startswith('【') or t.startswith('（') or t.startswith('参考') or t.startswith('行间距') or t.startswith('序号') or t.startswith('除期刊') or t.startswith('最新版') or t == '':
            to_clean.append(pp._element)
    for elem in to_clean:
        if elem.getparent() is body:
            body.remove(elem)
    if to_clean:
        print(f"  ✓ 清理参考文献后模板残留: {len(to_clean)}个段落")

    print(f"  ✓ 参考文献: {len(references)}条已填充")
    return doc


def phase4_handle_acknowledgment(doc, content, style_defs):
    print("阶段4: 处理致谢...")
    ack_text = content.get("acknowledgment", "")
    if not ack_text:
        print("  ⚠ 无致谢内容")
        return doc

    if isinstance(ack_text, list):
        ack_text = '\n'.join(ack_text)

    ack_idx = None
    ack_candidates = []
    for i, p in enumerate(doc.paragraphs):
        if "致    谢" in p.text and "目录" not in p.text:
            ack_candidates.append(i)
    if ack_candidates:
        ack_idx = ack_candidates[-1]

    if ack_idx is None:
        print("  ⚠ 未找到致谢标题")
        return doc

    ack_target_para = None
    for ci in range(ack_idx + 1, min(ack_idx + 30, len(doc.paragraphs))):
        ap = doc.paragraphs[ci]
        t = ap.text.strip()
        style_name = ap.style.name if ap.style else ""
        if "外文" in t or "附录" in t:
            break
        if style_name == "Title":
            break
        if t.startswith("×××") or t.startswith("感谢") or t.startswith("本论文") or t.startswith("衷心") or t.startswith("在"):
            ack_target_para = ap
            break
        if t == "" and len(ap.runs) <= 1:
            ack_target_para = ap
            break

    if ack_target_para is not None:
        ack_lines = [line for line in ack_text.split("\n") if line.strip()]
        body = doc.element.body
        ack_elem = ack_target_para._element

        first_para_text = ack_lines[0]
        merge_runs_to_first(ack_target_para)
        if ack_target_para.runs:
            ack_target_para.runs[0].text = first_para_text
        ack_target_para.style = doc.styles['正文——玄锐暮']
        for r in ack_target_para.runs[1:]:
            r.text = ""

        last_elem = ack_elem
        for line_text in ack_lines[1:]:
            elem = create_para_element("正文——玄锐暮", line_text, style_defs)
            last_elem.addnext(elem)
            last_elem = elem

        print("  ✓ 致谢内容已填充")
    else:
        print("  ⚠ 未找到致谢占位符")

    return doc


def phase5_handle_appendix(doc, content, style_defs):
    print("阶段5: 处理附录（已禁用——用户要求移除附录）...")
    print("  ✓ 附录已跳过")
    return doc

    appendix_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "附    录" in p.text and i > 100:
            appendix_idx = i
            break

    if appendix_idx is not None:
        body = doc.element.body
        to_clean = []
        to_clear_text = []
        for ci in range(appendix_idx + 1, len(doc.paragraphs)):
            ap = doc.paragraphs[ci]
            t = ap.text.strip()
            if "致    谢" in t or "外文原文" in t:
                break
            is_template_junk = (t.startswith("×××") or t.startswith("#include") or
                    t.startswith("void main") or t.startswith("{double") or
                    t.startswith("double V") or t.startswith("#define") or
                    "中国字" in t or t.startswith("程序清单") or
                    t.startswith("（1）字体") or t.startswith("（2）段落") or
                    t.startswith("（3）强调") or t.startswith("【附录"))
            if not is_template_junk:
                continue
            has_sectPr = ap._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
            if has_sectPr is not None:
                to_clear_text.append(ap._element)
            else:
                to_clean.append(ap._element)
        for elem in to_clean:
            if elem.getparent() is body:
                body.remove(elem)
        for elem in to_clear_text:
            for r in elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
                elem.remove(r)

        app_title_elem = doc.paragraphs[appendix_idx]._element
        last_elem = app_title_elem
        for line in appendix_lines:
            elem = create_para_element("程序清单-附录——玄锐暮", line, style_defs)
            last_elem.addnext(elem)
            last_elem = elem
        print(f"  ✓ 附录内容已添加: {len(appendix_lines)}行")
    else:
        foreign_idx = None
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip() == "外文原文及译文" and i > 100:
                foreign_idx = i
                break

        if foreign_idx is None:
            print("  ⚠ 未找到外文原文及译文标题，无法插入附录")
            return doc

        foreign_elem = doc.paragraphs[foreign_idx]._element

        app_title_elem = create_para_element(
            "致谢、外文原文及翻译——玄锐暮", "附    录", style_defs, page_break=True
        )
        foreign_elem.addprevious(app_title_elem)

        last_elem = app_title_elem
        for line in appendix_lines:
            elem = create_para_element("程序清单-附录——玄锐暮", line, style_defs)
            last_elem.addnext(elem)
            last_elem = elem
        print(f"  ✓ 附录内容已插入（致谢与外文原文之间）: {len(appendix_lines)}行")

    return doc


def phase6_handle_foreign_text(doc, content, style_defs):
    print("阶段6: 处理外文原文及译文...")
    foreign_original = content.get("foreign_original", {})
    foreign_translation = content.get("foreign_translation", {})

    has_original = foreign_original.get("title") or foreign_original.get("body")
    if not has_original:
        print("  ⚠ 无外文原文内容，跳过")
        return doc

    foreign_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "外文原文及译文" and i > 100:
            foreign_idx = i
            break

    if foreign_idx is None:
        print("  ⚠ 未找到外文原文及译文标题，跳过")
        return doc

    body = doc.element.body
    to_clean = []
    sectPr_preserved = 0
    for ci in range(foreign_idx + 1, len(doc.paragraphs)):
        ap = doc.paragraphs[ci]
        has_sectPr = ap._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
        if has_sectPr is not None:
            sectPr_preserved += 1
            continue
        to_clean.append(ap._element)

    for elem in to_clean:
        if elem.getparent() is body:
            body.remove(elem)
    if sectPr_preserved > 0:
        print(f"  ✓ 保护了 {sectPr_preserved} 个含sectPr的段落（页码/分节符保护）")

    foreign_title_para = doc.paragraphs[foreign_idx]
    last_elem = foreign_title_para._element

    if foreign_original.get("title"):
        elem = create_para_element("外文原文标题——玄锐暮", foreign_original["title"], style_defs)
        last_elem.addnext(elem)
        last_elem = elem

    if foreign_original.get("author_affiliation"):
        for line in foreign_original["author_affiliation"].split("\n"):
            line = line.strip()
            if line:
                elem = create_para_element("外文原文作者、单位——玄锐暮", line, style_defs)
                last_elem.addnext(elem)
                last_elem = elem

    if foreign_original.get("abstract"):
        abstract_text = foreign_original["abstract"]
        clean_check = abstract_text.replace("**", "").strip()
        if not clean_check.upper().startswith("ABSTRACT"):
            abstract_text = "ABSTRACT：" + abstract_text
        elem = create_para_element("外文原文摘要——玄锐暮", abstract_text, style_defs)
        last_elem.addnext(elem)
        last_elem = elem

    if foreign_original.get("keywords"):
        keywords_text = foreign_original["keywords"]
        clean_check = keywords_text.replace("**", "").strip()
        if not clean_check.upper().startswith("KEYWORDS"):
            keywords_text = "KEYWORDS：" + keywords_text
        keywords_text = keywords_text.replace("; ", "；").replace(";", "；").replace(": ", "：").replace(":", "：")
        elem = create_para_element("外文原文关键词——玄锐暮", keywords_text, style_defs)
        last_elem.addnext(elem)
        last_elem = elem

    for item in foreign_original.get("body", []):
        if item["type"] == "heading":
            elem = create_para_element("外文原文一二级标题——玄锐暮", item["text"], style_defs)
        elif item["type"] == "para":
            elem = create_para_element("外文原文正文——玄锐暮", item["text"], style_defs)
        else:
            continue
        last_elem.addnext(elem)
        last_elem = elem

    has_translation = foreign_translation.get("title") or foreign_translation.get("body")
    if has_translation:
        pb_elem = create_page_break_element()
        last_elem.addnext(pb_elem)
        last_elem = pb_elem

        if foreign_translation.get("title"):
            elem = create_para_element("翻译标题——玄锐暮", foreign_translation["title"], style_defs)
            last_elem.addnext(elem)
            last_elem = elem

        if foreign_translation.get("author_affiliation"):
            for line in foreign_translation["author_affiliation"].split("\n"):
                line = line.strip()
                if line:
                    elem = create_para_element("外文原文作者、单位——玄锐暮", line, style_defs)
                    last_elem.addnext(elem)
                    last_elem = elem

        if foreign_translation.get("abstract"):
            trans_abstract = foreign_translation["abstract"]
            clean_check = trans_abstract.replace("**", "").strip()
            if not clean_check.startswith("摘要"):
                trans_abstract = "摘要：" + trans_abstract
            elem = create_para_element("翻译摘要——玄锐暮", trans_abstract, style_defs)
            last_elem.addnext(elem)
            last_elem = elem

        if foreign_translation.get("keywords"):
            trans_keywords = foreign_translation["keywords"]
            clean_check = trans_keywords.replace("**", "").strip()
            if not clean_check.startswith("关键词"):
                trans_keywords = "关键词：" + trans_keywords
            elem = create_para_element("翻译关键词——玄锐暮", trans_keywords, style_defs)
            last_elem.addnext(elem)
            last_elem = elem

        for item in foreign_translation.get("body", []):
            if item["type"] == "heading":
                elem = create_para_element("翻译一二级标题——玄锐暮", item["text"], style_defs)
            elif item["type"] == "para":
                elem = create_para_element("正文——玄锐暮", item["text"], style_defs)
            else:
                continue
            last_elem.addnext(elem)
            last_elem = elem

    print(f"  ✓ 外文原文及译文已添加（结构化样式）")
    return doc


def verify_output(output_path, content):
    print("\n阶段7: 生成后验证...")
    errors = []

    doc = Document(output_path)
    template_doc = Document(TEMPLATE_PATH)

    print("  验证1: 5页保护验证...")
    protected_keywords = ["原创性声明", "版权使用授权书", "学术诚信承诺书"]

    cover_end = 0
    for i, p in enumerate(template_doc.paragraphs):
        if p.style and "玄锐暮" in (p.style.name or ""):
            cover_end = i
            break

    abstract_cn_start = None
    abstract_en_start = None
    for i, p in enumerate(template_doc.paragraphs):
        if p.text.strip() == "摘    要" and abstract_cn_start is None:
            abstract_cn_start = i
        elif p.text.strip() == "ABSTRACT" and abstract_en_start is None:
            abstract_en_start = i

    if cover_end > 0:
        for i in range(min(cover_end, len(doc.paragraphs), len(template_doc.paragraphs))):
            if abstract_cn_start is not None and abstract_en_start is not None:
                if abstract_cn_start <= i <= abstract_en_start:
                    continue
                if abstract_en_start < i < cover_end:
                    continue
            if doc.paragraphs[i].text != template_doc.paragraphs[i].text:
                errors.append(f"封面保护失败: 第{i}段被修改 '{template_doc.paragraphs[i].text[:30]}...' → '{doc.paragraphs[i].text[:30]}...'")
                break

    toc_start = None
    toc_end = None
    template_toc_start = None
    template_toc_end = None

    for i, p in enumerate(doc.paragraphs):
        if toc_start is None and p.text.strip().replace(" ", "") == "目录":
            toc_start = i
        elif toc_start is not None and toc_end is None and "第一章" in p.text:
            toc_end = i
            break

    for i, p in enumerate(template_doc.paragraphs):
        if template_toc_start is None and p.text.strip().replace(" ", "") == "目录":
            template_toc_start = i
        elif template_toc_start is not None and template_toc_end is None and "第一章" in p.text:
            template_toc_end = i
            break

    if all(v is not None for v in [toc_start, toc_end, template_toc_start, template_toc_end]):
        if (toc_end - toc_start) != (template_toc_end - template_toc_start):
            errors.append(f"目录保护失败: 目录段落数不一致 (输出{toc_end - toc_start}段 vs 模板{template_toc_end - template_toc_start}段)")
        else:
            for i in range(toc_end - toc_start):
                if doc.paragraphs[toc_start + i].text != template_doc.paragraphs[template_toc_start + i].text:
                    errors.append(f"目录保护失败: 目录第{i}段被修改 '{template_doc.paragraphs[template_toc_start + i].text[:30]}...' → '{doc.paragraphs[toc_start + i].text[:30]}...'")
                    break

    for i, (op, tp) in enumerate(zip(doc.paragraphs, template_doc.paragraphs)):
        if any(kw in tp.text for kw in protected_keywords):
            in_protected = False
            for j in range(max(0, i - 2), min(len(template_doc.paragraphs), i + 30)):
                if any(kw in template_doc.paragraphs[j].text for kw in protected_keywords):
                    in_protected = True
                    break
            if in_protected and op.text != tp.text:
                errors.append(f"5页保护失败: 第{i}段被修改 '{tp.text[:30]}...' → '{op.text[:30]}...'")
    print(f"  ✓ 5页保护验证完成")

    print("  验证2: 强制替换验证...")
    abstract_cn = content.get("abstract_cn", "")
    if abstract_cn:
        found_abstract = False
        for p in doc.paragraphs:
            if len(p.text) > 50 and abstract_cn[:30] in p.text:
                found_abstract = True
                break
        if not found_abstract:
            errors.append("强制替换失败: 中文摘要未被替换")
    chapters = content.get("chapters", [])
    if chapters:
        found_ch1 = False
        for p in doc.paragraphs:
            if "第一章" in p.text:
                found_ch1 = True
                break
        if not found_ch1:
            errors.append("强制替换失败: 正文未被替换")
    print(f"  ✓ 强制替换验证完成")

    print("  验证3: 框架完整性验证...")
    required_frameworks = [
        "摘    要", "ABSTRACT", "目    录",
        "第一章", "第二章", "第三章", "第四章", "第五章", "第六章",
        "参考文献", "致    谢", "外文原文"
    ]
    found_texts = set()
    for p in doc.paragraphs:
        for fw in required_frameworks:
            if fw in p.text:
                found_texts.add(fw)
    for fw in required_frameworks:
        if fw not in found_texts:
            errors.append(f"框架缺失: '{fw}' 未找到")
    print(f"  ✓ 框架完整性验证完成 (找到 {len(found_texts)}/{len(required_frameworks)})")

    print("  验证4: 引用编号顺序验证...")
    ref_pattern = re.compile(r'\[(\d+)\]')
    first_appearance = {}
    max_ref_num = 0
    ch1_start = None
    for i, p in enumerate(doc.paragraphs):
        if "第一章" in p.text and ch1_start is None:
            ch1_start = i
            break

    if ch1_start is not None:
        for i in range(ch1_start, len(doc.paragraphs)):
            pt = doc.paragraphs[i].text
            if "参考文献" in pt and i > ch1_start:
                break
            for m in ref_pattern.finditer(pt):
                num = int(m.group(1))
                if num > max_ref_num:
                    max_ref_num = num
                if num not in first_appearance:
                    first_appearance[num] = i

    if first_appearance:
        sorted_by_appearance = sorted(first_appearance.items(), key=lambda x: x[1])
        appearance_order = [num for num, _ in sorted_by_appearance]
        expected_order = sorted(appearance_order)
        if appearance_order != expected_order:
            disorder_examples = []
            for idx in range(len(appearance_order) - 1):
                if appearance_order[idx] > appearance_order[idx + 1]:
                    disorder_examples.append(f"[{appearance_order[idx]}]出现在[{appearance_order[idx+1]}]之前")
            errors.append(f"引用编号顺序错误! 正文中引用出现顺序为{appearance_order}，但应为{expected_order}。示例：{'; '.join(disorder_examples[:3])}")
        else:
            print(f"    ✓ 引用编号按首次出现顺序排列: {appearance_order}")
    print(f"  ✓ 引用编号顺序验证完成")

    if errors:
        print(f"\n  ❌ 验证失败! 发现 {len(errors)} 个问题:")
        for e in errors:
            print(f"    - {e}")
        print("  ⚠ 请修复上述问题后重新生成!")
    else:
        print("  ✅ 所有验证通过!")


def main():
    if not os.path.exists(TEMPLATE_PATH):
        print(f"错误: 找不到模板文件 {TEMPLATE_PATH}")
        sys.exit(1)

    md_path = sys.argv[1] if len(sys.argv) > 1 else CONTENT_MD_PATH
    output_path = sys.argv[2] if len(sys.argv) > 2 else DEFAULT_OUTPUT

    if not os.path.exists(md_path):
        print(f"错误: 找不到内容文件 {md_path}")
        print("请先由AI Agent生成thesis_content.md文件")
        sys.exit(1)

    content = parse_markdown(md_path)
    content = _fix_reference_numbering(content)

    print("=" * 60)
    print("  仁爱学院论文生成工具 (v4 双重保险)")
    print("=" * 60)
    print(f"\n模板: {TEMPLATE_PATH}")
    print(f"内容: {md_path}")
    print(f"输出: {output_path}\n")

    style_defs = read_style_defs(Document(TEMPLATE_PATH))
    print(f"  ✓ 从模板读取 {len(style_defs)} 个玄锐暮样式定义（双重保险）")

    temp_path = phase1_render_template(content)
    doc = phase2_add_chapters(temp_path, content, style_defs)
    doc = phase3_handle_references(doc, content, style_defs)
    doc = phase4_handle_acknowledgment(doc, content, style_defs)
    doc = phase5_handle_appendix(doc, content, style_defs)
    doc = phase6_handle_foreign_text(doc, content, style_defs)

    doc.save(output_path)

    if os.path.exists(temp_path):
        os.remove(temp_path)
        print(f"  ✓ 清理临时文件")

    verify_output(output_path, content)

    print(f"\n{'=' * 60}")
    print(f"  论文已生成: {output_path}")
    print(f"{'=' * 60}")
    print("\n⚠ 请注意以下手动操作：")
    print("  1. 在标注位置插入图片（架构图、流程图、E-R图等）")
    print("  2. 确认交叉引用可点击跳转（generate.py已自动插入REF字段和书签）")
    print("  3. 更新目录（右键目录→更新域→更新整个目录）")
    print("  4. 上传维普自检（查重+AIGC检测+格式检测）")


if __name__ == "__main__":
    main()
